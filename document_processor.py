import os
import re
from typing import List, Optional

def parse_document(file_path: str) -> str:
    """
    Parse different document types (PDF, DOCX, TXT) and extract text content.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        str: Extracted text from the document
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return parse_pdf(file_path)
    elif file_extension == '.docx':
        return parse_docx(file_path)
    elif file_extension == '.txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def parse_pdf(file_path: str) -> str:
    """
    Extract text from PDF files.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
    """
    try:
        import PyPDF2
        
        extracted_text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                extracted_text += page.extract_text() + "\n\n"
        
        # Clean up text: remove excessive whitespace
        cleaned_text = re.sub(r'\s+', ' ', extracted_text).strip()
        return cleaned_text
    
    except ImportError:
        raise ImportError("Unable to process PDF files. Please install PyPDF2.")

def parse_docx(file_path: str) -> str:
    """
    Extract text from DOCX files.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    try:
        import docx
        
        document = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    except ImportError:
        raise ImportError("Unable to process DOCX files. Please install python-docx.")

def parse_txt(file_path: str) -> str:
    """
    Extract text from plain text files.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        str: Contents of the text file
    """
    with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
        return file.read()

def chunk_text(text: str, tokenizer=None, max_chunk_size: int = 1000) -> List[str]:
    """
    Split text into chunks based on character count.
    
    Args:
        text: Text to split into chunks
        tokenizer: Not used in this version, kept for API compatibility
        max_chunk_size: Maximum number of characters per chunk
        
    Returns:
        List[str]: List of text chunks
    """
    # Split text into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = []
    current_chunk_length = 0
    
    for sentence in sentences:
        sentence_length = len(sentence)
        
        # If a single sentence is too long, split it further
        if sentence_length > max_chunk_size:
            # Add current chunk if not empty
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_chunk_length = 0
            
            # Split long sentence into smaller pieces
            words = sentence.split()
            temp_piece = []
            temp_length = 0
            
            for word in words:
                word_length = len(word) + 1  # +1 for space
                
                if temp_length + word_length <= max_chunk_size:
                    temp_piece.append(word)
                    temp_length += word_length
                else:
                    chunks.append(' '.join(temp_piece))
                    temp_piece = [word]
                    temp_length = word_length
            
            if temp_piece:
                chunks.append(' '.join(temp_piece))
        
        # Normal case: sentence fits within character limit
        elif current_chunk_length + sentence_length <= max_chunk_size:
            current_chunk.append(sentence)
            current_chunk_length += sentence_length
        else:
            # Finish current chunk and start a new one
            chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_chunk_length = sentence_length
    
    # Add the last chunk if not empty
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks
